"""Final report assembly.

Produces two artifacts on the state:
  - state["report_data"]: structured dict (stable keys, machine-consumable)
  - state["report"]:      rendered plain-text report for humans

An optional .docx rendering is gated behind REFASSIST_WRITE_DOCX_REPORT
(concurrent runs would race on a shared file, and the API builds its own
per-request document from report_data).
"""
import os
import re
import time
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from ..state import PipelineState
from ..tools.utils import (
    authors_to_list, safe_str, format_doi_link, normalize_text,
    token_similarity, isbn_valid, DOI_SYNTAX_RE,
)

# Fields the verification agents can actually confirm against sources —
# everything else in the output is taken from the input or a single source.
_VERIFIABLE = {"title", "authors", "journal_name", "journal_abbrev", "year",
               "month", "volume", "issue", "pages", "doi"}

EXPORTS_DIR = os.path.join(os.getcwd(), "exports")

SRC_LABELS = {
    "doi-agreement": "DOI agreement",
    "crossref-exact": "Crossref (exact match)",
    "consensus": "Consensus",
    "crossref": "Crossref",
    "ieeexplore": "IEEE Xplore",
    "openalex": "OpenAlex",
    "semanticscholar": "Semantic Scholar",
    "pubmed": "PubMed",
    "arxiv": "arXiv",
    "openlibrary": "Open Library",
    "dblp": "DBLP",
    "datacite": "DataCite",
    "europepmc": "Europe PMC",
    "unpaywall": "Unpaywall",
    "biorxiv": "bioRxiv/medRxiv",
    "doaj": "DOAJ",
    "googlebooks": "Google Books",
    "extracted": "Provided input",
    "normalize": "Normalization",
    "verify/llm": "Verification agent",
    "authoritative": "Authoritative source",
    "nlm": "NLM Catalog",
    "candidates": "Candidate records",
    "": "Unknown",
}

# Pseudo-values pushed into `corrections` by verify_journal_abbrev — these are
# diagnostics, not applied corrections, and belong in the warnings section.
_NON_CORRECTIONS = {"missing journal name", "journal not found", "not found"}

_FIELD_ORDER = [
    "title", "authors", "editors", "journal_name", "journal_abbrev",
    "verified_journal_abbrev", "conference_name", "book_title", "volume",
    "issue", "pages", "year", "month", "doi", "publisher", "location",
    "edition", "isbn", "url",
]


def _src_label(code: str) -> str:
    return SRC_LABELS.get((code or "").lower(), code or "Unknown")


def _fmt_value(field: str, val: Any) -> str:
    if field in ("authors", "editors"):
        return ", ".join(authors_to_list(val))
    return safe_str(val)


def _collect_evidence(state: PipelineState) -> List[Tuple[str, str]]:
    """(label, url) links supporting the SELECTED record, deduped by URL.

    Only candidates that back the chosen match count as evidence — links from
    unrelated search results would misattribute the correction.
    """
    seen = set()
    lines: List[Tuple[str, str]] = []

    def add(label: str, url: str):
        url = normalize_text(url)
        if url and url not in seen:
            seen.add(url)
            lines.append((label, url))

    best = state.get("best", {}) or {}
    best_doi = normalize_text(best.get("doi") or "").lower().replace("doi:", "")
    best_title = normalize_text(best.get("title") or "").lower()
    if best.get("doi"):
        add("DOI", format_doi_link(best["doi"]))

    def _supports_best(c: dict) -> bool:
        if not best:
            return False
        c_doi = normalize_text(c.get("doi") or "").lower().replace("doi:", "")
        if best_doi and c_doi:
            return c_doi == best_doi
        c_title = normalize_text(c.get("title") or "").lower()
        return bool(best_title) and c_title == best_title

    for c in state.get("candidates", []) or []:
        if not _supports_best(c):
            continue
        src = (c.get("source") or "").lower()
        raw = c.get("raw") or {}

        if src == "crossref":
            if raw.get("DOI"):
                add("Crossref (DOI)", format_doi_link(raw["DOI"]))
            if raw.get("URL"):
                add("Publisher (from Crossref)", raw["URL"])
        elif src == "ieeexplore":
            add("IEEE Xplore", raw.get("html_url") or "")
            if raw.get("doi"):
                add("DOI", format_doi_link(raw["doi"]))
        elif src == "openalex":
            add("OpenAlex", raw.get("id") or "")
            if raw.get("doi"):
                add("DOI", format_doi_link(raw["doi"]))
        elif src == "semanticscholar":
            eid = raw.get("externalIds") or {}
            doi = normalize_text(eid.get("DOI") or raw.get("doi") or "")
            if doi:
                add("Semantic Scholar", f"https://www.semanticscholar.org/doi/{doi}")
            elif raw.get("paperId"):
                add("Semantic Scholar", f"https://www.semanticscholar.org/paper/{raw['paperId']}")
        elif src == "pubmed":
            pmid = normalize_text(raw.get("uid") or "")
            if pmid.isdigit():
                add("PubMed", f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/")
        elif src == "arxiv":
            aid = normalize_text((state.get("extracted", {}) or {}).get("arxiv_id") or "")
            if aid:
                add("arXiv", f"https://arxiv.org/abs/{aid}")

        if c.get("doi"):
            add("DOI", format_doi_link(c["doi"]))

    return lines


def _split_corrections(raw_changes: List[Tuple[str, Any, Any]]):
    """Separate real applied corrections from diagnostic pseudo-entries."""
    corrections, notes = [], []
    for f, old, new in raw_changes or []:
        new_s = safe_str(new)
        if new_s.lower() in _NON_CORRECTIONS or new_s.lower().startswith("verification error"):
            old_s = safe_str(old)
            suffix = f" (was: '{old_s}')" if old_s else ""
            notes.append(f"Journal abbreviation not in NLM Catalog{suffix} — "
                         "expected for non-biomedical journals; the abbreviation "
                         "shown comes from the metadata sources instead.")
            continue
        corrections.append((f, old, new))
    return corrections, notes


def _warnings(state: PipelineState) -> List[str]:
    warn: List[str] = []
    ex = state.get("extracted", {}) or {}
    rtype = normalize_text(state.get("type") or "")

    y = safe_str(ex.get("year"))
    if y and not re.fullmatch(r"(18|19|20)\d{2}", y):
        warn.append(f"Suspicious year value: '{y}'")
    if rtype in ("journal article", "conference paper") and not normalize_text(ex.get("doi")):
        warn.append("No DOI found for an article/conference reference")
    pages = normalize_text(ex.get("pages"))
    if pages:
        nums = re.findall(r"\d+", pages.replace("—", "-").replace("–", "-"))
        if "-" in pages and len(nums) >= 2 and nums[0] == nums[1]:
            warn.append(f"Page range '{pages}' starts and ends on the same page")
        elif pages.isdigit():
            warn.append(f"Single page '{pages}' — verify whether it should be a range")

    doi = normalize_text(ex.get("doi")).replace("doi:", "").strip()
    if doi and not DOI_SYNTAX_RE.fullmatch(doi):
        warn.append(f"DOI '{doi}' does not look like a valid DOI")
    isbn = normalize_text(ex.get("isbn"))
    if isbn and not isbn_valid(isbn):
        warn.append(f"ISBN '{isbn}' fails checksum validation")

    alts = state.get("version_alternatives") or []
    if alts:
        warn.append(
            f"This work appears in multiple published forms ({' and '.join(alts)}) — "
            "verify you are citing the intended version.")
    return warn


def _sources_log(state: PipelineState) -> List[Dict[str, str]]:
    """Per-source consultation log: an author should see exactly which
    databases were asked and what each returned."""
    candidates = state.get("candidates", []) or []
    by_src: Dict[str, list] = {}
    for c in candidates:
        by_src.setdefault((c.get("source") or "").lower(), []).append(c)
    timed_out = {k.split("|")[0] for k in (state.get("_timed_out_jobs") or set())}

    ex = state.get("extracted", {}) or {}
    rtype = (state.get("type") or "").lower()
    book_like = rtype in ("book", "book chapter") or bool(
        ex.get("publisher") and not ex.get("journal_name"))
    has_doi = bool(normalize_text(ex.get("doi")))

    doi_only = {"datacite": "DOI-registry lookup only",
                "unpaywall": "per-DOI service only",
                "biorxiv": "bioRxiv/medRxiv DOIs only"}
    log = []
    for s in state.get("_sources") or []:
        name = getattr(s, "NAME", "").lower()
        enabled_check = getattr(s, "_enabled", None)
        if enabled_check is not None and not enabled_check():
            outcome = "not enabled (no API key configured)"
        elif name in ("openlibrary", "googlebooks") and not book_like:
            outcome = "not queried (book references only)"
        elif name in doi_only and not has_doi:
            outcome = f"not queried ({doi_only[name]}; no DOI resolved)"
        elif name in by_src:
            hits = by_src[name]
            vias = sorted({c.get("_via") or "?" for c in hits})
            outcome = f"{len(hits)} candidate record(s) via {', '.join(vias)} search"
        elif name in timed_out:
            outcome = "timed out — no answer within the per-source limit"
        else:
            outcome = "queried — no matching records"
        log.append({"source": _src_label(name), "outcome": outcome})
    return log


def _field_disagreements(state: PipelineState) -> List[str]:
    """Where sources that matched this work disagreed on a field, disclose
    every value seen, who said it, and which value was used."""
    ex = state.get("extracted", {}) or {}
    final_title = normalize_text(ex.get("title") or "")
    if not final_title:
        return []

    def norm(fld: str, v: Any) -> str:
        s = normalize_text(str(v or "")).lower().replace("doi:", "")
        if fld == "doi":  # OpenAlex returns DOIs as full URLs
            s = s.removeprefix("https://doi.org/").removeprefix("http://doi.org/")
        return s.replace("–", "-").replace("—", "-") if fld == "pages" else s

    relevant = [c for c in state.get("candidates", []) or []
                if token_similarity(final_title, normalize_text(c.get("title") or "")) >= 0.85]
    out = []
    for fld in ("year", "volume", "issue", "pages", "doi"):
        seen: Dict[str, set] = {}
        for c in relevant:
            v = norm(fld, c.get(fld))
            if v:
                seen.setdefault(v, set()).add(_src_label(c.get("source")))
        if len(seen) > 1:
            chosen = norm(fld, ex.get(fld))
            alts = " vs ".join(f"'{v}' ({', '.join(sorted(srcs))})"
                               for v, srcs in sorted(seen.items()))
            used = f" — used '{chosen}'" if chosen else " — field left unset"
            out.append(f"{fld}: {alts}{used}")
    return out


def _action_items(state: PipelineState, status: str, retracted: bool,
                  rw_info: Dict[str, Any], failed: List[str]) -> List[str]:
    """What the author must do before submitting, in order of severity."""
    if status == "rejected":
        return ["Input was not recognized as a bibliographic reference — check for "
                "missing fields or formatting problems and resubmit."]
    items = []
    fab = state.get("fabrication") or {}
    # The author-mismatch signal gets its own (more detailed) item below.
    fab_signals = [s for s in fab.get("signals", []) if s.get("code") != "authors_disjoint"]
    if fab.get("risk") == "high" and fab_signals:
        items.append(
            "LIKELY FABRICATED REFERENCE: "
            + "; ".join(s["detail"] for s in fab_signals)
            + ". Do not cite this unless you can produce the original document.")
    elif fab.get("risk") == "medium" and fab_signals:
        items.append(
            "Possible fabricated reference: "
            + "; ".join(s["detail"] for s in fab_signals)
            + ". Verify against the original source before citing.")
    mismatch = state.get("author_mismatch") or {}
    if mismatch:
        pub_auths = ", ".join(mismatch.get("published_authors", [])[:6]) or "unknown"
        loc = f" (DOI: {mismatch['doi']})" if mismatch.get("doi") else ""
        items.append(
            "POSSIBLE FABRICATED CITATION: a published work matches this title "
            f"almost exactly, but its authors are entirely different — published "
            f"authors: {pub_auths}{loc}. AI-generated citations often invent "
            "authors for real titles. Verify against the original source before citing.")
        if mismatch.get("retracted"):
            items.append("Additionally, that published work is RETRACTED — if it is "
                         "the work you meant to cite, do not cite it as valid research.")
    if retracted:
        notice = (rw_info or {}).get("notice_doi")
        items.append("Do NOT cite this as valid research — the work is retracted"
                     + (f" (notice: https://doi.org/{notice})" if notice else "")
                     + ". If citing the retraction itself, say so explicitly.")
    elif rw_info:
        items.append(f"A '{rw_info.get('nature', 'notice')}' is on record for this work — "
                     "review it before citing.")
    if state.get("version_alternatives"):
        items.append("Multiple published versions exist ("
                     + " and ".join(state["version_alternatives"])
                     + ") — confirm the version you cite is the one you used.")
    if status == "unverified":
        items.append("No trustworthy online match was found — verify every field "
                     "against the original source before submitting.")
    elif failed:
        items.append("Manually verify: " + ", ".join(sorted(failed))
                     + " (not confirmed by any consulted source).")
    return items


def _processing_meta(state: PipelineState) -> Dict[str, Any]:
    """Engine provenance: what produced this result, with what, in how long."""
    llm = state.get("_llm")
    cfg = state.get("_cfg")
    provider = getattr(llm, "provider", "none")
    if provider == "azure":
        from ..llms.adapter import azure_chat_deployment
        engine = f"azure:{azure_chat_deployment()}"
    elif provider == "dummy":
        engine = "heuristics only (no LLM configured)"
    else:
        model = {"openai": getattr(cfg, "openai_model", ""),
                 "anthropic": getattr(cfg, "anthropic_model", ""),
                 "ollama": getattr(cfg, "ollama_model", "")}.get(provider, "")
        engine = f"{provider}:{model}" if model else provider

    started = state.get("_started_at")
    try:
        from ..tools.retractionwatch import get_rw_db
        db = get_rw_db()
        rw_status = f"active ({db.count:,} records)" if db.ready else "not loaded this run"
    except Exception:
        rw_status = "unavailable"

    return {
        "duration_s": round(time.time() - started, 1) if started else None,
        "lookup_rounds": state.get("hops", 0),
        "correction_rounds": state.get("attempts", 0),
        "engine": engine,
        "retraction_db": rw_status,
    }


def _status_and_confidence(state: PipelineState, corrections: List) -> Tuple[str, str, str]:
    """Returns (status, status_label, confidence)."""
    if state.get("_skip_pipeline"):
        return "rejected", "Rejected — input is not a bibliographic reference", "n/a"

    fab = state.get("fabrication") or {}
    if fab.get("risk") == "high":
        return ("suspect",
                "Suspect — likely fabricated or unlocatable; do not cite without "
                "verifying the original document exists",
                "n/a")

    best = state.get("best", {}) or {}
    prov = state.get("provenance", {}) or {}
    ver = {k: v for k, v in (state.get("verification") or {}).items() if k != "is_reference"}
    ver_ratio = (sum(1 for v in ver.values() if v) / len(ver)) if ver else 0.0

    trusted_doi = normalize_text(prov.get("doi", "")).lower() in {"doi-agreement", "crossref-exact"} \
        or any(normalize_text(v).lower() == "crossref-exact" for v in prov.values())

    if not best or (best.get("source") == "extracted"):
        # No candidates at all, or every candidate failed the author gate and
        # select_best fell back to the input's own fields — nothing was
        # actually verified online.
        return ("unverified",
                "Unverified — no trustworthy online match; formatted from provided data",
                "low")

    if trusted_doi and ver_ratio >= 0.75:
        confidence = "high"
    elif ver_ratio >= 0.6:
        confidence = "medium"
    else:
        confidence = "low"

    # Calibration: a merge that CONTRADICTS data the user supplied is where
    # wrong-work merges hide — it must never present as fully trusted.
    if state.get("identity_conflict"):
        return ("corrected",
                "Verified with conflicts — the matched record contradicts "
                "identifiers you supplied; review before use",
                "low")
    ex0 = state.get("_original_extracted") or {}
    contradicted = sum(
        1 for f, old, _new in corrections
        if f in ("volume", "issue", "pages", "doi", "isbn", "authors", "title")
        and safe_str(old).strip() and safe_str(ex0.get(f, "")).strip()
    )
    if contradicted >= 2 and confidence == "high":
        confidence = "medium"

    if corrections:
        return "corrected", "Verified against online sources — corrections applied", confidence
    return "verified", "Verified against online sources — no corrections needed", confidence


def _build_report_data(state: PipelineState) -> Dict[str, Any]:
    ex = state.get("extracted", {}) or {}
    best = state.get("best", {}) or {}
    prov = state.get("provenance", {}) or {}
    audit = state.get("audit", {}) or {}
    ver = state.get("verification", {}) or {}

    corrections, notes = _split_corrections(state.get("corrections", []) or [])
    status, status_label, confidence = _status_and_confidence(state, corrections)

    ver_checks = {k: bool(v) for k, v in ver.items() if k != "is_reference"}

    # verified: True (confirmed by sources) / False (checked, not confirmed) /
    # None (field type the verification agents cannot check)
    fields = []
    for f in _FIELD_ORDER:
        val = ex.get(f) if ex.get(f) not in (None, "", []) else best.get(f)
        if val in (None, "", []):
            continue
        verified = ver_checks.get(f) if f in _VERIFIABLE and status != "unverified" else None
        fields.append({
            "field": f,
            "value": _fmt_value(f, val),
            "source": _src_label(audit.get(f) or prov.get(f) or ("extracted" if ex.get(f) else "")),
            "verified": verified,
        })
    retracted = bool(state.get("retracted") or best.get("retracted"))
    rw_info = state.get("retraction_info") or {}
    warnings = _warnings(state) if status != "rejected" else []
    if retracted:
        msg = ("This work has been RETRACTED by its publisher — "
               "do not cite it without noting the retraction.")
        if rw_info.get("reasons"):
            msg += f" Reasons (Retraction Watch): {'; '.join(rw_info['reasons'][:4])}."
        warnings.insert(0, msg)
    elif rw_info:
        # Expression of concern / correction: advisory, not a retraction.
        nature = rw_info.get("nature") or "notice"
        detail = f" ({'; '.join(rw_info['reasons'][:3])})" if rw_info.get("reasons") else ""
        warnings.insert(0, f"{nature} on record"
                           f"{' since ' + rw_info['date'] if rw_info.get('date') else ''}"
                           f" per Retraction Watch{detail} — review before citing.")
    mismatch = state.get("author_mismatch") or {}
    if mismatch:
        warnings.insert(0, "Cited authors do not match any published work with this "
                           "title — possible fabricated or misattributed citation.")
    fab = state.get("fabrication") or {}
    if fab.get("risk") in ("high", "medium") and fab.get("signals"):
        head = "LIKELY FABRICATED" if fab["risk"] == "high" else "Possibly fabricated"
        warnings.insert(0, f"{head} — " + "; ".join(s["detail"] for s in fab["signals"]))
    warnings.extend(notes)

    failed = [f["field"] for f in fields if f["verified"] is False]
    unconfirmed = [f["field"] for f in fields
                   if f["verified"] is None and f["source"] == "Provided input"]

    try:
        from .. import __version__ as _ver
    except Exception:
        _ver = "unknown"

    return {
        "action_items": _action_items(state, status, retracted, rw_info, failed),
        "sources_consulted": _sources_log(state) if status != "rejected" else [],
        "disagreements": _field_disagreements(state) if status != "rejected" else [],
        "unconfirmed_fields": unconfirmed,
        "processing": _processing_meta(state),
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "pipeline_version": _ver,
        "input": state.get("reference", ""),
        "status": status,
        "status_label": status_label,
        "confidence": confidence,
        "retracted": retracted,
        "retraction_info": rw_info or None,
        "author_mismatch": mismatch or None,
        "fabrication": fab or None,
        "type": state.get("type") or ("n/a" if status == "rejected" else "other"),
        "doi": safe_str(best.get("doi") or ex.get("doi")),
        "formatted": state.get("formatted", "") or "",
        # The formatter that ACTUALLY produced the output sets this flag;
        # inferring it from side artifacts previously mislabeled the strategy.
        "formatting_strategy": (
            "n/a" if status == "rejected"
            else state.get("_formatter") or "Rule-based IEEE formatter (deterministic)"
        ),
        "corrections": [
            {
                "field": f,
                "old": _fmt_value(f, old) or "(missing)",
                "new": _fmt_value(f, new) or "(removed)",
                "source": _src_label(audit.get(f) or prov.get(f) or ""),
            }
            for f, old, new in corrections
        ],
        "fields": fields,
        "verification": ver_checks,
        "evidence": [{"label": l, "url": u} for l, u in _collect_evidence(state)],
        "sources_searched": sorted({
            _src_label(c.get("source")) for c in state.get("candidates", []) or [] if c.get("source")
        }),
        "warnings": warnings,
        "message": state.get("verification_message", ""),
        "fingerprint": state.get("_fp", ""),
    }


def _render_text(d: Dict[str, Any]) -> str:
    W = 78
    bar = "=" * W
    thin = "-" * W
    lines: List[str] = []

    lines += [bar, "REFASSIST REFERENCE REPORT".center(W), bar]
    lines.append(f"Generated : {d['generated_at']}   (refassist {d['pipeline_version']})")
    lines.append(f"Status    : {d['status_label']}")
    if d["confidence"] != "n/a":
        lines.append(f"Confidence: {d['confidence'].capitalize()}")
    lines.append(f"Type      : {d['type'].title() if d['type'] else 'Unknown'}")
    if d["doi"]:
        lines.append(f"DOI       : {d['doi']}")

    if d.get("retracted"):
        lines += [
            "",
            "!" * W,
            "RETRACTION NOTICE".center(W),
            "This work has been RETRACTED by its publisher.".center(W),
            "It should not be cited as valid research without noting the retraction.".center(W),
        ]
        info = d.get("retraction_info") or {}
        if info.get("date"):
            lines.append(f"Retracted: {info['date']} (Retraction Watch)".center(W))
        if info.get("reasons"):
            lines.append(f"Reasons: {'; '.join(info['reasons'][:4])}".center(W))
        if info.get("notice_doi"):
            lines.append(f"Retraction notice: https://doi.org/{info['notice_doi']}".center(W))
        lines.append("!" * W)

    fab = d.get("fabrication") or {}
    if (fab.get("risk") or "none") in ("high", "medium") and fab.get("signals"):
        lines += ["", "!" * W, "FABRICATION CHECK".center(W),
                  f"Risk: {fab['risk'].upper()}".center(W), "-" * W]
        for s in fab["signals"]:
            lines.append(f"  - {s['detail']}")
        lines.append("!" * W)

    # What the author must do, before anything else
    lines += ["", "ACTION REQUIRED" if d["action_items"] else "ACTION REQUIRED: none", thin]
    if d["action_items"]:
        for i, a in enumerate(d["action_items"], 1):
            lines.append(f"  {i}. {a}")
    else:
        lines.append("  This reference can be used as shown below.")

    lines += ["", "INPUT (as submitted)", thin, f"  {d['input']}"]

    if d["status"] == "rejected":
        lines += ["", "NOTE", thin, f"  {d['message'] or 'Input was not processed.'}", bar]
        return "\n".join(lines)

    lines += ["", "FINAL REFERENCE (IEEE)", thin,
              f"  {d['formatted'] or '(no formatted output produced)'}",
              f"  Produced by: {d['formatting_strategy']}"]

    lines += ["", f"CORRECTIONS APPLIED ({len(d['corrections'])})", thin]
    if d["corrections"]:
        for c in d["corrections"]:
            lines.append(f"  - {c['field']}: {c['old']}  ->  {c['new']}   [{c['source']}]")
    else:
        lines.append("  None — the reference matched authoritative records as provided.")

    # Field table with per-field verification marks
    lines += ["", "FIELDS  (v = confirmed by sources, ! = could not be confirmed,",
              "         · = informational only — not independently checkable)", thin]
    width = max((len(f["field"]) for f in d["fields"]), default=0)
    for f in d["fields"]:
        mark = {True: "v", False: "!"}.get(f["verified"], "·")
        lines.append(f"  {mark} {f['field']:<{width}} : {f['value']}   ({f['source']})")
    if d["unconfirmed_fields"]:
        lines.append(f"  Taken from your input without online confirmation: "
                     f"{', '.join(d['unconfirmed_fields'])}")

    # Full transparency: every database consulted and what it said
    lines += ["", "SOURCES CONSULTED", thin]
    for s in d["sources_consulted"]:
        lines.append(f"  - {s['source']:<18} {s['outcome']}")

    if d["disagreements"]:
        lines += ["", "SOURCE DISAGREEMENTS (all values seen, and which was used)", thin]
        for g in d["disagreements"]:
            lines.append(f"  - {g}")

    lines += ["", "EVIDENCE (verify these yourself)", thin]
    if d["evidence"]:
        for e in d["evidence"]:
            lines.append(f"  - {e['label']}: {e['url']}")
    else:
        lines.append("  No online evidence captured.")

    if d["warnings"]:
        lines += ["", "WARNINGS", thin]
        for w in d["warnings"]:
            lines.append(f"  - {w}")

    p = d["processing"]
    lines += ["", "PROCESSING LOG", thin,
              f"  Engine            : {p['engine']}",
              f"  Retraction check  : Crossref + OpenAlex + PubMed + Retraction Watch ({p['retraction_db']})",
              f"  Lookup rounds     : {p['lookup_rounds']}   Correction rounds: {p['correction_rounds']}",
              (f"  Processing time   : {p['duration_s']}s" if p["duration_s"] is not None else "  Processing time   : n/a"),
              f"  Pipeline          : refassist {d['pipeline_version']}",
              f"  Fingerprint       : {d['fingerprint']}",
              bar]
    return "\n".join(lines)


def build_report(state: PipelineState) -> PipelineState:
    data = _build_report_data(state)
    state["report_data"] = data
    state["report"] = _render_text(data)

    # -------- Word report (opt-in; see module docstring) --------
    if os.getenv("REFASSIST_WRITE_DOCX_REPORT", "").lower() not in ("1", "true", "yes"):
        return state

    try:
        from docx import Document
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        doc = Document()
        doc.add_heading("RefAssist Reference Report", level=1)
        doc.add_paragraph(f"Generated {data['generated_at']} — refassist {data['pipeline_version']}")
        doc.add_paragraph(f"{data['status_label']} (confidence: {data['confidence']})")
        doc.add_heading("Input", level=2)
        doc.add_paragraph(data["input"])
        doc.add_heading("Final Reference (IEEE)", level=2)
        doc.add_paragraph(data["formatted"] or "(none)")
        if data["corrections"]:
            doc.add_heading("Corrections", level=2)
            t = doc.add_table(rows=1, cols=4)
            hdr = t.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Field", "Before", "After", "Source"
            for c in data["corrections"]:
                row = t.add_row().cells
                row[0].text, row[1].text, row[2].text, row[3].text = c["field"], c["old"], c["new"], c["source"]
        doc.add_heading("Fields", level=2)
        for f in data["fields"]:
            doc.add_paragraph(f"{f['field']}: {f['value']}  ({f['source']})")
        if data["warnings"]:
            doc.add_heading("Warnings", level=2)
            for w in data["warnings"]:
                doc.add_paragraph(w, style="List Bullet")
        doc.save(os.path.join(EXPORTS_DIR, "report.docx"))
        state["report_path"] = os.path.join(EXPORTS_DIR, "report.docx")
    except Exception:
        ...  # report file generation must never fail the pipeline

    return state
