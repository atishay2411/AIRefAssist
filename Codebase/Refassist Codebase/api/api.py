from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from refassist.graphs import run_one
from refassist.config import PipelineConfig
from docx import Document
from typing import Optional
import tempfile, os, zipfile, uuid, re

app = FastAPI(title="RefAssist API", version="0.4.0")

# Static + templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


# ---------- Helper: robust reference splitting ----------
def split_references(text: str) -> list[str]:
    """
    Split references by common patterns:
    - Numbered: 1. Ref ...
    - IEEE style: [1] Ref ...
    - Bullets: - Ref ... or • Ref ...
    - Or fallback: newlines
    """
    text = text.strip()
    if not text:
        return []

    # Normalize line breaks
    text = re.sub(r"\r\n?", "\n", text)

    # Split using regex patterns
    refs = re.split(r"(?:\n+|(?=\[\d+\])|(?=\d+\.)|(?=•)|(?=- ))", text)

    # Cleanup extra symbols and whitespace
    refs = [r.strip(" \n-•") for r in refs if r.strip(" \n-•")]
    return refs


# ---------- Helper: process references ----------
async def process_references(refs: list[str]):
    formatted_refs = []
    report_doc = Document()
    report_doc.add_heading("Reference Processing Report", 0)

    for idx, ref in enumerate(refs, start=1):
        try:
            out = await run_one(ref.strip(), PipelineConfig())
            formatted = out.get("formatted", ref)
            formatted_refs.append(formatted)

            # Add report entry
            report_doc.add_heading(f"Reference {idx}", level=2)
            report_doc.add_paragraph(f"Original: {ref}")
            report_doc.add_paragraph(f"Processed: {formatted}")
            report_doc.add_paragraph(out.get("report", "No changes"))
        except Exception as e:
            formatted_refs.append(ref + " [ERROR]")
            report_doc.add_heading(f"Reference {idx}", level=2)
            report_doc.add_paragraph(f"Original: {ref}")
            report_doc.add_paragraph(f"Error: {str(e)}")

    return formatted_refs, report_doc


# ---------- UI route ----------
@app.get("/", response_class=HTMLResponse)
async def get_home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# ---------- API: handle copy-paste OR file ----------
@app.post("/v1/upload")
async def upload_references(
    references: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    refs_text = None

    # Case 1: pasted references
    if references and references.strip():
        refs_text = references.strip()

    # Case 2: uploaded file
    elif file:
        content = await file.read()
        refs_text = content.decode("utf-8", errors="ignore")

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    # Use robust splitter
    refs = split_references(refs_text)
    formatted_refs, report_doc = await process_references(refs)

    # Save results into temp folder
    tmpdir = tempfile.mkdtemp()
    txt_path = os.path.join(tmpdir, "formatted_references.txt")
    docx_path = os.path.join(tmpdir, "report.docx")
    zip_path = os.path.join(tmpdir, f"results_{uuid.uuid4().hex}.zip")

    # Write formatted refs as numbered IEEE-style list
    with open(txt_path, "w", encoding="utf-8") as f:
        for idx, ref in enumerate(formatted_refs, start=1):
            f.write(f"[{idx}] {ref}\n")

    # Save Word report
    report_doc.save(docx_path)

    # Zip the files
    with zipfile.ZipFile(zip_path, "w") as zipf:
        zipf.write(txt_path, "formatted_references.txt")
        zipf.write(docx_path, "report.docx")

    return FileResponse(zip_path, filename="results.zip", media_type="application/zip")
