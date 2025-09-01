# from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
# from fastapi.responses import HTMLResponse, StreamingResponse
# from fastapi.staticfiles import StaticFiles
# from fastapi.templating import Jinja2Templates
# from refassist.graphs import run_one
# from refassist.config import PipelineConfig
# from docx import Document
# from typing import Optional, List, Tuple
# import re
# import asyncio
# import io
# import zipfile
# import logging

# # ---------- Setup ----------
# app = FastAPI(title="RefAssist API", version="0.6.0")
# app.mount("/static", StaticFiles(directory="static"), name="static")
# templates = Jinja2Templates(directory="templates")

# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger("refassist")

# def split_references(text: str) -> List[str]:
#     """
#     Robust reference splitter:
#     - Handles IEEE [1], numbered 1., bullets -, •
#     - Keeps multi-line references together
#     """
#     text = text.strip()
#     if not text:
#         return []

#     # Normalize line breaks
#     text = re.sub(r"\r\n?", "\n", text)
#     lines = text.split("\n")

#     refs = []
#     current_ref = []

#     # Patterns for reference start
#     patterns = [
#         re.compile(r"^\s*\[\d+\]"),  # [1] style
#         re.compile(r"^\s*\d+\."),    # 1. style
#         re.compile(r"^\s*[-•]"),     # bullet style
#     ]

#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue  # skip empty lines

#         # Check if this line starts a new reference
#         if any(p.match(line) for p in patterns):
#             if current_ref:
#                 refs.append(" ".join(current_ref).strip())
#             current_ref = [re.sub(r"^[-•\[\]\d\.]\s*", "", line)]  # remove marker
#         else:
#             # continuation of previous reference
#             current_ref.append(line)

#     if current_ref:
#         refs.append(" ".join(current_ref).strip())

#     return refs


# # ---------- Helper: process references concurrently ----------
# async def process_references(refs: List[str]) -> Tuple[List[str], Document]:
#     formatted_refs: List[str] = []
#     report_doc = Document()
#     report_doc.add_heading("Reference Processing Report", 0)

#     async def process_single(idx: int, ref: str) -> Tuple[str, dict]:
#         try:
#             out = await run_one(ref.strip(), PipelineConfig())
#             formatted = out.get("formatted", ref)
#             report_entry = {
#                 "idx": idx,
#                 "original": ref,
#                 "formatted": formatted,
#                 "report": out.get("report", "No changes"),
#                 "matching_fields": out.get("matching_fields", [])  # NEW: Include matching fields
#             }
#             return formatted, report_entry
#         except Exception as e:
#             logger.exception(f"Error processing reference {idx}")
#             return ref + " [ERROR]", {
#                 "idx": idx,
#                 "original": ref,
#                 "formatted": None,
#                 "report": f"Error: {str(e)}",
#                 "matching_fields": []
#             }

#     tasks = [process_single(i + 1, ref) for i, ref in enumerate(refs)]
#     results = await asyncio.gather(*tasks, return_exceptions=False)

#     for formatted, entry in results:
#         formatted_refs.append(formatted)
#         report_doc.add_heading(f"Reference {entry['idx']}", level=2)
#         report_doc.add_paragraph(f"Original: {entry['original']}")
#         if entry['formatted']:
#             report_doc.add_paragraph(f"Processed: {entry['formatted']}")
#         report_doc.add_paragraph(f"Matching Fields: {', '.join(entry['matching_fields']) if entry['matching_fields'] else 'None'}")
#         report_doc.add_paragraph(entry['report'])

#     return formatted_refs, report_doc


# # ---------- UI route ----------
# @app.get("/", response_class=HTMLResponse)
# async def get_home(request: Request):
#     return templates.TemplateResponse("index.html", {"request": request})


# # ---------- API: handle copy-paste OR file ----------
# @app.post("/v1/upload")
# async def upload_references(
#     references: Optional[str] = Form(None),
#     file: Optional[UploadFile] = File(None)
# ):
#     refs_text: Optional[str] = None

#     # Case 1: pasted references
#     if references and references.strip():
#         refs_text = references.strip()

#     # Case 2: uploaded file
#     elif file:
#         content = await file.read()
#         try:
#             if file.filename.lower().endswith(".docx"):
#                 from io import BytesIO
#                 doc = Document(BytesIO(content))
#                 refs_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
#             else:
#                 refs_text = content.decode("utf-8", errors="ignore")
#         except Exception as e:
#             logger.exception("Failed to read uploaded file")
#             raise HTTPException(status_code=400, detail=f"Invalid file format: {str(e)}")

#     if not refs_text:
#         raise HTTPException(status_code=400, detail="No references provided")

#     # Split and process references
#     refs = split_references(refs_text)
#     if not refs:
#         raise HTTPException(status_code=400, detail="No references detected in input")

#     formatted_refs, report_doc = await process_references(refs)

#     # Create in-memory ZIP
#     zip_buffer = io.BytesIO()
#     with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
#         # Write TXT
#         txt_content = "\n".join(f"[{i+1}] {ref}" for i, ref in enumerate(formatted_refs))
#         zipf.writestr("formatted_references.txt", txt_content)

#         # Write DOCX report
#         docx_buffer = io.BytesIO()
#         report_doc.save(docx_buffer)
#         docx_buffer.seek(0)
#         zipf.writestr("report.docx", docx_buffer.read())

#     zip_buffer.seek(0)
#     return StreamingResponse(
#         zip_buffer,
#         media_type="application/zip",
#         headers={"Content-Disposition": 'attachment; filename="results.zip"'}
#     )

from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from refassist.graphs import run_one
from refassist.config import PipelineConfig
from docx import Document
from typing import Optional, List, Tuple
import re
import asyncio
import io
import zipfile
import logging

# ---------- Setup ----------
app = FastAPI(title="RefAssist API", version="0.6.0")
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("refassist")

def split_references(text: str) -> List[str]:
    """
    Robust reference splitter:
    - Handles IEEE [1], numbered 1., bullets -, •
    - Keeps multi-line references together
    """
    text = text.strip()
    if not text:
        return []

    # Normalize line breaks
    text = re.sub(r"\r\n?", "\n", text)
    lines = text.split("\n")

    refs = []
    current_ref = []

    # Patterns for reference start
    patterns = [
        re.compile(r"^\s*\[\d+\]"),  # [1] style
        re.compile(r"^\s*\d+\."),    # 1. style
        re.compile(r"^\s*[-•]"),     # bullet style
    ]

    for line in lines:
        line = line.strip()
        if not line:
            continue  # skip empty lines

        # Check if this line starts a new reference
        if any(p.match(line) for p in patterns):
            if current_ref:
                refs.append(" ".join(current_ref).strip())
            current_ref = [re.sub(r"^[-•\[\]\d\.]\s*", "", line)]  # remove marker
        else:
            # continuation of previous reference
            current_ref.append(line)

    if current_ref:
        refs.append(" ".join(current_ref).strip())

    return refs


# ---------- Helper: process references concurrently ----------
async def process_references(refs: List[str]) -> Tuple[List[str], Document]:
    formatted_refs: List[str] = []
    report_doc = Document()
    report_doc.add_heading("Reference Processing Report", 0)

    async def process_single(idx: int, ref: str) -> Tuple[str, dict]:
        try:
            out = await run_one(ref.strip(), PipelineConfig())
            formatted = out.get("formatted", ref)
            report_entry = {
                "idx": idx,
                "original": ref,
                "formatted": formatted,
                "report": out.get("report", "No changes")
            }
            return formatted, report_entry
        except Exception as e:
            logger.exception(f"Error processing reference {idx}")
            return ref + " [ERROR]", {
                "idx": idx,
                "original": ref,
                "formatted": None,
                "report": f"Error: {str(e)}"
            }

    tasks = [process_single(i + 1, ref) for i, ref in enumerate(refs)]
    results = await asyncio.gather(*tasks, return_exceptions=False)

    for formatted, entry in results:
        formatted_refs.append(formatted)
        report_doc.add_heading(f"Reference {entry['idx']}", level=2)
        report_doc.add_paragraph(f"Original: {entry['original']}")
        if entry['formatted']:
            report_doc.add_paragraph(f"Processed: {entry['formatted']}")
        report_doc.add_paragraph(entry['report'])

    return formatted_refs, report_doc


# ---------- OLD UI route (DISABLED) ----------
# @app.get("/", response_class=HTMLResponse)
# async def get_home(request: Request):
#     return templates.TemplateResponse("index.html", {"request": request})

# ---------- NEW UI route ----------
@app.get("/", response_class=HTMLResponse)
async def get_new_home(request: Request):
    from fastapi.responses import FileResponse
    return FileResponse("new_UI/index.html")

# Mount new UI static files
app.mount("/new_ui", StaticFiles(directory="new_UI"), name="new_ui")


# ---------- NEW API: process references for new UI ----------
@app.post("/api/process")
async def process_references_api(
    references: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    """
    Process references and return structured data for the new UI
    """
    refs_text: Optional[str] = None

    # Case 1: pasted references
    if references and references.strip():
        refs_text = references.strip()

    # Case 2: uploaded file
    elif file:
        content = await file.read()
        try:
            if file.filename.lower().endswith(".docx"):
                from io import BytesIO
                doc = Document(BytesIO(content))
                refs_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                refs_text = content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.exception("Failed to read uploaded file")
            raise HTTPException(status_code=400, detail=f"Invalid file format: {str(e)}")

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    # Split and process references
    refs = split_references(refs_text)
    if not refs:
        raise HTTPException(status_code=400, detail="No references detected in input")

    # Process references and collect detailed results
    formatted_refs: List[str] = []
    detailed_reports: List[dict] = []
    
    async def process_single_detailed(idx: int, ref: str) -> Tuple[str, dict]:
        try:
            out = await run_one(ref.strip(), PipelineConfig())
            formatted = out.get("formatted", ref)
            report_entry = {
                "idx": idx,
                "original": ref,
                "formatted": formatted,
                "report": out.get("report", "No changes"),
                "status": "success"
            }
            return formatted, report_entry
        except Exception as e:
            logger.exception(f"Error processing reference {idx}")
            error_formatted = ref + " [ERROR]"
            report_entry = {
                "idx": idx,
                "original": ref,
                "formatted": error_formatted,
                "report": f"Error: {str(e)}",
                "status": "error"
            }
            return error_formatted, report_entry

    tasks = [process_single_detailed(i + 1, ref) for i, ref in enumerate(refs)]
    results = await asyncio.gather(*tasks, return_exceptions=False)

    for formatted, entry in results:
        formatted_refs.append(formatted)
        detailed_reports.append(entry)

    # Generate formatted output text
    formatted_output = "\n".join(f"[{i+1}] {ref}" for i, ref in enumerate(formatted_refs))
    
    # Generate detailed report text
    report_lines = ["Reference Processing Report", "=" * 30, ""]
    
    # Summary
    total_refs = len(detailed_reports)
    success_count = sum(1 for r in detailed_reports if r["status"] == "success")
    error_count = total_refs - success_count
    
    report_lines.append(f"Total references processed: {total_refs}")
    report_lines.append(f"Successfully processed: {success_count}")
    report_lines.append(f"Errors encountered: {error_count}")
    report_lines.append("")
    
    # Individual reference details
    for entry in detailed_reports:
        report_lines.append(f"Reference {entry['idx']}:")
        report_lines.append(f"Original: {entry['original']}")
        if entry['status'] == 'success':
            report_lines.append(f"Formatted: {entry['formatted']}")
        report_lines.append(f"Notes: {entry['report']}")
        report_lines.append("")
    
    detailed_report_text = "\n".join(report_lines)
    
    # Generate preview (first few lines of detailed report)
    preview_lines = report_lines[:15]  # First 15 lines as preview
    if len(report_lines) > 15:
        preview_lines.append("... (see full report for complete analysis)")
    preview_text = "\n".join(preview_lines)

    return {
        "success": True,
        "total_references": total_refs,
        "formatted_output": formatted_output,
        "detailed_report": detailed_report_text,
        "preview": preview_text,
        "summary": {
            "total": total_refs,
            "success": success_count,
            "errors": error_count
        }
    }


# ---------- API: download full report ----------
@app.post("/api/download-report")
async def download_full_report(
    references: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    """
    Generate and download full report as ZIP file
    """
    # Reuse the existing logic from the original upload endpoint
    refs_text: Optional[str] = None

    if references and references.strip():
        refs_text = references.strip()
    elif file:
        content = await file.read()
        try:
            if file.filename.lower().endswith(".docx"):
                from io import BytesIO
                doc = Document(BytesIO(content))
                refs_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                refs_text = content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.exception("Failed to read uploaded file")
            raise HTTPException(status_code=400, detail=f"Invalid file format: {str(e)}")

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    refs = split_references(refs_text)
    if not refs:
        raise HTTPException(status_code=400, detail="No references detected in input")

    formatted_refs, report_doc = await process_references(refs)

    # Create in-memory ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        # Write TXT
        txt_content = "\n".join(f"[{i+1}] {ref}" for i, ref in enumerate(formatted_refs))
        zipf.writestr("formatted_references.txt", txt_content)

        # Write DOCX report
        docx_buffer = io.BytesIO()
        report_doc.save(docx_buffer)
        docx_buffer.seek(0)
        zipf.writestr("report.docx", docx_buffer.read())

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="refassist_report.zip"'}
    )


# ---------- API: handle copy-paste OR file (LEGACY) ----------
@app.post("/v1/upload")
async def upload_references(
    references: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    refs_text: Optional[str] = None

    # Case 1: pasted references
    if references and references.strip():
        refs_text = references.strip()

    # Case 2: uploaded file
    elif file:
        content = await file.read()
        try:
            if file.filename.lower().endswith(".docx"):
                from io import BytesIO
                doc = Document(BytesIO(content))
                refs_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                refs_text = content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.exception("Failed to read uploaded file")
            raise HTTPException(status_code=400, detail=f"Invalid file format: {str(e)}")

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    # Split and process references
    refs = split_references(refs_text)
    if not refs:
        raise HTTPException(status_code=400, detail="No references detected in input")

    formatted_refs, report_doc = await process_references(refs)

    # Create in-memory ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        # Write TXT
        txt_content = "\n".join(f"[{i+1}] {ref}" for i, ref in enumerate(formatted_refs))
        zipf.writestr("formatted_references.txt", txt_content)

        # Write DOCX report
        docx_buffer = io.BytesIO()
        report_doc.save(docx_buffer)
        docx_buffer.seek(0)
        zipf.writestr("report.docx", docx_buffer.read())

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="results.zip"'}
    )
