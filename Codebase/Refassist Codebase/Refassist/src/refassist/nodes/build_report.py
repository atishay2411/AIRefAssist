import os
import re
from typing import Dict, List, Tuple, Any
from docx import Document
from docx.shared import Pt
from ..state import PipelineState
from ..tools.utils import authors_to_list, safe_str, format_doi_link, normalize_month_field, normalize_text

TEMPLATE_PATH = os.path.join(os.getcwd(), "Template.docx")
EXPORTS_DIR = os.path.join(os.getcwd(), "exports")

SRC_LABELS = {
    "doi-agreement": "DOI agreement",
    "consensus": "Consensus",
    "crossref": "Crossref",
    "ieeexplore": "IEEE Xplore",
    "openalex": "OpenAlex",
    "semanticscholar": "Semantic Scholar",
    "pubmed": "PubMed",
    "arxiv": "arXiv",
    "normalize": "Normalization",
    "verify/llm": "LLM suggestion",
    "nlm": "NLM Catalog",
    "": "Unknown"
}

def _src_label(code: str) -> str:
    return SRC_LABELS.get((code or "").lower(), code or "Unknown")

def _fmt_list(xs: List[str]) -> str:
    return ", ".join(sorted(set([x for x in xs if x])))

def _doi_link(doi: str) -> str:
    return format_doi_link(doi)

def _collect_evidence(state: PipelineState) -> List[Tuple[str, str]]:
    """
    Build a list of (Source, URL) evidence lines from candidates and best.
    Dedupes by URL.
    """
    seen = set()
    lines: List[Tuple[str,str]] = []

    best = state.get("best", {}) or {}
    bd = normalize_text(best.get("doi") or "")
    if bd:
        url = _doi_link(bd)
        if url and url not in seen:
            seen.add(url); lines.append(("DOI", url))

    # Candidates often carry rich "raw" records we can turn into links.
    for c in state.get("candidates", []) or []:
        src = (c.get("source") or "").lower()
        raw = c.get("raw") or {}

        def add(label: str, url: str):
            if url and url not in seen:
                seen.add(url); lines.append((label, url))

        if src == "crossref":
            doi = normalize_text(raw.get("DOI") or "")
            add("Crossref (DOI)", _doi_link(doi))
            # Some Crossref records include a 'URL' to the publisher page:
            if raw.get("URL"):
                add("Publisher (from Crossref)", normalize_text(raw.get("URL")))

        elif src == "ieeexplore":
            # IEEE Xplore API fields are 'html_url'/'pdf_url'
            add("IEEE Xplore", normalize_text(raw.get("html_url") or ""))
            if raw.get("pdf_url"):
                add("IEEE Xplore PDF", normalize_text(raw.get("pdf_url")))
            doi = normalize_text(raw.get("doi") or "")
            add("DOI", _doi_link(doi))

        elif src == "openalex":
            # OpenAlex 'id' is a URL; prefer that; also include DOI if present
            add("OpenAlex", normalize_text(raw.get("id") or ""))
            doi = normalize_text(raw.get("doi") or "")
            add("DOI", _doi_link(doi))

        elif src == "semanticscholar":
            # Prefer the canonical S2 UI DOI page when available
            eid = raw.get("externalIds") or {}
            doi = normalize_text(eid.get("DOI") or raw.get("doi") or "")
            if doi:
                add("Semantic Scholar (DOI)", f"https://www.semanticscholar.org/doi/{doi}")
            # Fallback: paperId
            pid = normalize_text(raw.get("paperId") or "")
            if pid:
                add("Semantic Scholar", f"https://www.semanticscholar.org/paper/{pid}")

        elif src == "pubmed":
            # esummary record has a numeric uid/pmid
            pmid = normalize_text(raw.get("uid") or "")
            if pmid.isdigit():
                add("PubMed", f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/")

        elif src == "arxiv":
            # If we captured an arXiv id in extracted or best
            ex = state.get("extracted", {}) or {}
            aid = normalize_text(ex.get("arxiv_id") or "")
            if not aid:
                aid = normalize_text(best.get("arxiv_id") or "")
            if aid:
                add("arXiv", f"https://arxiv.org/abs/{aid}")

        # Any explicit URL field present in normalized cand
        url_norm = normalize_text(c.get("url") or "")
        if url_norm:
            add(_src_label(src), url_norm)

        # As a final fallback, add DOI from normalized candidate
        doi_norm = normalize_text(c.get("doi") or "")
        if doi_norm:
            add("DOI", _doi_link(doi_norm))

    return lines

def _format_corrections(changes: List[Tuple[str, Any, Any]], prov: Dict[str,str], audit: Dict[str,str]) -> str:
    if not changes:
        return "No corrections were needed."

    lines = []
    for f, old, new in changes:
        prov_source = _src_label(audit.get(f) or prov.get(f) or "Unknown")
        if f == "authors":
            old_str = ", ".join(authors_to_list(old)) if old else "MISSING"
            new_str = ", ".join(authors_to_list(new)) if new else "MISSING"
        else:
            old_str = safe_str(old) if old else "MISSING"
            new_str = safe_str(new) if new else "MISSING"
        lines.append(f"- {f}: {old_str} → {new_str}  (by: {prov_source})")
    return "\n".join(lines)

def _format_provenance(best: Dict[str,Any], ex: Dict[str,Any], prov: Dict[str,str], audit: Dict[str,str]) -> str:
    lines = []
    all_fields = [
        "title", "authors", "journal_name", "journal_abbrev", "verified_journal_abbrev",
        "conference_name", "volume", "issue", "pages", "year", "month", "doi",
        "publisher", "location", "edition", "isbn", "url"
    ]
    for f in all_fields:
        val = best.get(f) or ex.get(f)
        if val is None or val == "":
            continue
        if f == "authors":
            val = ", ".join(authors_to_list(val))
        source_code = prov.get(f) or audit.get(f) or ""
        src = _src_label(source_code)
        lines.append(f"- {f}: {val}  (source: {src})")
    return "\n".join(lines)

def _trust_summary(state: PipelineState) -> str:
    prov = state.get("provenance", {}) or {}
    best = state.get("best", {}) or {}
    cands = state.get("candidates", []) or []

    reason = "No trusted online match"
    trusted = bool(best)
    if trusted:
        if normalize_text(prov.get("doi","")) == "doi-agreement":
            reason = "DOI agreement across sources"
        else:
            # infer strong-match rationale
            used = [p for p in set(prov.values()) if p]
            reason = f"Strict title/author/venue match from {_fmt_list([_src_label(u) for u in used])}"

    sources_used = _fmt_list([_src_label(s.get("source")) for s in cands])
    return f"Trusted online match: {'Yes' if trusted else 'No'}\nMatch rationale: {reason}\nSources searched: {sources_used or 'None'}"

def _nlm_note(state: PipelineState) -> str:
    ex = state.get("extracted", {}) or {}
    v = ex.get("verified_journal_abbrev") or ""
    if v:
        return f"NLM Catalog ISO Abbrev: {v}"
    return "NLM Catalog ISO Abbrev: (not verified or not applicable)"

def _warnings(state: PipelineState) -> List[str]:
    warn: List[str] = []
    ex = state.get("extracted", {}) or {}
    rtype = normalize_text(state.get("type") or "")

    # Year plausibility
    y = safe_str(ex.get("year"))
    if y and (not re.fullmatch(r"(18|19|20)\d{2}", y)):
        warn.append(f"Suspicious year value: '{y}'")
    # Missing DOI for journal/conference material
    if rtype in ("journal article","conference paper"):
        if not normalize_text(ex.get("doi")):
            warn.append("Missing DOI for an article/conference reference")
    # Single page vs range
    pages = normalize_text(ex.get("pages"))
    if pages:
        nums = re.findall(r"\d+", pages.replace("—","-").replace("–","-"))
        if "-" in pages and len(nums) >= 2 and nums[0] == nums[1]:
            warn.append(f"Detected fake page range '{pages}' collapsed to single page")
        elif pages.isdigit():
            warn.append(f"Single page '{pages}' — verify whether it should be a range")
    return warn

def build_report(state: PipelineState) -> PipelineState:
    ex = state.get("extracted", {}) or {}
    best = state.get("best", {}) or {}
    prov = state.get("provenance", {}) or {}
    audit = state.get("audit", {}) or {}
    changes = state.get("corrections", []) or []
    ver = state.get("verification", {}) or {}
    matching_fields = set(state.get("matching_fields", []) or [])
    formatted = state.get("formatted", "") or ""
    failed = [k for k, v in (ver or {}).items() if not v and k != "is_reference"]

    # Overview
    overview = (
        f"- Type detected: {state.get('type','Unknown')}\n"
        f"- DOI: {best.get('doi') or ex.get('doi') or 'Not available'}\n"
        f"- Primary source for DOI: {_src_label(prov.get('doi') or 'Unknown')}\n"
        f"- {_trust_summary(state)}"
    )

    # Field verification snapshot
    verification = (
        f"Fields matched authoritative sources: {', '.join(sorted(matching_fields)) or 'None'}\n"
        f"Fields needing attention: {', '.join(sorted(failed)) or 'None'}"
    )

    # Corrections & provenance
    corrections = _format_corrections(changes, prov, audit)
    provenance = _format_provenance(best, ex, prov, audit)

    formatting = "LLM-based formatting applied successfully " if formatted else "LLM formatting failed or skipped \nFalling back to rule-based IEEE formatter "

    final_reference = formatted or state.get("ieee_formatted", "") or "Error: No formatted reference available."

    # Evidence links
    evidence_lines = _collect_evidence(state)
    if evidence_lines:
        evidence_txt = "\n".join([f"- {label}: {url}" for (label, url) in evidence_lines])
    else:
        evidence_txt = "No online evidence URLs captured."

    # Journal abbreviations (NLM)
    nlm_txt = _nlm_note(state)

    # Warnings / anomalies
    warnings = _warnings(state)
    warnings_txt = "\n".join([f"- {w}" for w in warnings]) if warnings else "- None"

    # Reproducibility fingerprint
    fp = state.get("_fp") or ""

    # Put it all together
    state["report"] = f"""
IEEE Reference Report

1. Overview
{overview}

2. Field Verification
{verification}

3. Corrections Applied
{corrections}

4. Provenance (Source per Field)
{provenance}

5. Online Evidence (links)
{evidence_txt}

6. Journal Abbreviation Check
{nlm_txt}

7. Formatting Strategy
{formatting}

8. Final Formatted Reference
{final_reference}

9. Data Quality Warnings
{warnings_txt}

10. Reproducibility
- Fingerprint: {fp}
"""

    # -------- Word report (optional) --------
    if not os.path.exists(EXPORTS_DIR):
        os.makedirs(EXPORTS_DIR)

    try:
        if not os.path.exists(TEMPLATE_PATH):
            # Fallback basic .docx
            doc = Document()
            doc.add_heading("IEEE Reference Report", level=1)
            for section, block in [
                ("Overview", overview),
                ("Field Verification", verification),
                ("Corrections Applied", corrections),
                ("Provenance (Source per Field)", provenance),
                ("Online Evidence (links)", evidence_txt),
                ("Journal Abbreviation Check", nlm_txt),
                ("Formatting Strategy", formatting),
                ("Final Formatted Reference", final_reference),
                ("Data Quality Warnings", warnings_txt),
                ("Reproducibility", f"Fingerprint: {fp}"),
            ]:
                doc.add_heading(section, level=2)
                for line in block.split("\n"):
                    doc.add_paragraph(line)
            report_path = os.path.join(EXPORTS_DIR, "report.docx")
            doc.save(report_path)
            state["report_path"] = report_path
            return state

        # If a template exists, we replace placeholders if present,
        # and append new sections at the end so no info is lost.
        doc = Document(TEMPLATE_PATH)
        placeholders = {
            "{OVERVIEW}": overview,
            "{VERIFICATION}": verification,
            "{CORRECTIONS}": corrections,
            "{PROVENANCE}": provenance,
            "{EVIDENCE}": evidence_txt,
            "{NLM}": nlm_txt,
            "{FORMATTING}": formatting,
            "{FINAL_REFERENCE}": final_reference,
            "{WARNINGS}": warnings_txt,
            "{FINGERPRINT}": fp,
        }

        matched_any = False
        for p in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if placeholder in inline[i].text:
                            inline[i].text = inline[i].text.replace(placeholder, value)
                            matched_any = True

        if not matched_any:
            # Append content if template didn't contain placeholders
            doc.add_heading("IEEE Reference Report", level=1)
            for section, block in [
                ("Overview", overview),
                ("Field Verification", verification),
                ("Corrections Applied", corrections),
                ("Provenance (Source per Field)", provenance),
                ("Online Evidence (links)", evidence_txt),
                ("Journal Abbreviation Check", nlm_txt),
                ("Formatting Strategy", formatting),
                ("Final Formatted Reference", final_reference),
                ("Data Quality Warnings", warnings_txt),
                ("Reproducibility", f"Fingerprint: {fp}"),
            ]:
                doc.add_heading(section, level=2)
                for line in block.split("\n"):
                    doc.add_paragraph(line)

        report_path = os.path.join(EXPORTS_DIR, "report.docx")
        doc.save(report_path)
        state["report_path"] = report_path
    except Exception:
        # Fail silently for report file generation; JSON report still returned by API
        ...

    return state
