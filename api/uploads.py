"""Hardened text extraction for uploaded reference files.

Defenses, in the order they run per file:
  1. Extension allowlist and count limit.
  2. Chunked reads that abort the moment the per-file size limit is crossed
     (nothing unbounded is ever buffered).
  3. Magic-byte validation — the content must match the claimed type
     (%PDF header, DOCX zip structure, no NUL bytes in text formats).
  4. Decompression caps: DOCX zip-bomb guard, PDF page cap.
  5. CPU-bound parsers (pdfminer, python-docx) run in a worker thread so
     they never block the event loop.
  6. A total-extracted-text cap protects the downstream pipeline.

LaTeX/.bbl files get bibliography-aware cleanup (\\bibitem → paragraph
boundaries, command stripping) so the reference splitter sees clean text.
"""
import asyncio
import io
import logging
import os
import re
import time
import zipfile
from typing import List

from fastapi import HTTPException, UploadFile

logger = logging.getLogger("refassist.uploads")

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False


def env_pos_int(name: str, default: int, minimum: int = 1) -> int:
    """Positive int from env; malformed or too-small values fall back loudly
    instead of crashing startup or creating zero-capacity limits."""
    raw = os.getenv(name)
    try:
        v = int(raw) if raw is not None else default
    except (TypeError, ValueError):
        logger.warning("Invalid %s=%r; using default %d", name, raw, default)
        return default
    if v < minimum:
        logger.warning("%s=%d is below minimum %d; clamping", name, v, minimum)
        return minimum
    return v


MAX_UPLOAD_BYTES = env_pos_int("REFASSIST_MAX_UPLOAD_BYTES", 15 * 1024 * 1024, minimum=1024)
MAX_FILES = env_pos_int("REFASSIST_MAX_FILES", 10)
MAX_TOTAL_TEXT_CHARS = env_pos_int("REFASSIST_MAX_TEXT_CHARS", 2_000_000, minimum=10_000)
MAX_PDF_PAGES = env_pos_int("REFASSIST_MAX_PDF_PAGES", 200)
MAX_DOCX_UNCOMPRESSED = env_pos_int("REFASSIST_MAX_DOCX_UNCOMPRESSED", 100 * 1024 * 1024,
                                    minimum=1024 * 1024)

ALLOWED_EXTS = (".pdf", ".docx", ".tex", ".bbl", ".txt")

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x1f\x7f]")


def safe_name(filename) -> str:
    """Display-safe filename: basename only, control chars stripped, capped."""
    name = os.path.basename(str(filename or "upload"))
    name = _CONTROL_CHARS_RE.sub("", name).strip() or "upload"
    return name[:120]


def _ext(name: str) -> str:
    return "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""


async def read_limited(up: UploadFile, name: str) -> bytes:
    """Read an upload in chunks, aborting as soon as the size limit is crossed."""
    chunks: List[bytes] = []
    size = 0
    while True:
        chunk = await up.read(1024 * 1024)
        if not chunk:
            break
        size += len(chunk)
        if size > MAX_UPLOAD_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"'{name}' exceeds the {MAX_UPLOAD_BYTES // (1024*1024)} MB per-file limit."
            )
        chunks.append(chunk)
    return b"".join(chunks)


# ---------------------------------------------------------------------------
# Content validation
# ---------------------------------------------------------------------------

def validate_magic(ext: str, raw: bytes, name: str) -> None:
    """The bytes must match the claimed type — extensions are user input."""
    if ext == ".pdf":
        if b"%PDF-" not in raw[:1024]:
            raise HTTPException(400, f"'{name}' is not a valid PDF file.")
    elif ext == ".docx":
        if raw[:4] != b"PK\x03\x04":
            raise HTTPException(400, f"'{name}' is not a valid .docx file.")
        try:
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                names = z.namelist()
                if "[Content_Types].xml" not in names:
                    raise HTTPException(400, f"'{name}' is not a valid .docx file.")
                total = sum(i.file_size for i in z.infolist())
                if total > MAX_DOCX_UNCOMPRESSED:
                    raise HTTPException(
                        413, f"'{name}' decompresses to an unreasonable size and was rejected.")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(400, f"'{name}' is not a valid .docx file.")
    else:  # text formats
        if b"\x00" in raw[:8192]:
            raise HTTPException(400, f"'{name}' contains binary data, not text.")


def _decode_text(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", "replace")


# ---------------------------------------------------------------------------
# Format-specific extraction (sync parsers run via asyncio.to_thread)
# ---------------------------------------------------------------------------

_BIBITEM_RE = re.compile(r"\\bibitem(?:\[[^\]]*\])?\{[^}]*\}")
_TEX_COMMENT_RE = re.compile(r"(?m)(?<!\\)%.*$")
_TEX_HREF_RE = re.compile(r"\\href\{([^{}]*)\}\{([^{}]*)\}")
_TEX_WRAP_RE = re.compile(r"\\(?:emph|textit|textbf|texttt|textsc|textrm|text|url|path|mbox)\{([^{}]*)\}")
_TEX_CMD_RE = re.compile(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?")


def tex_to_text(s: str) -> str:
    """Bibliography-aware LaTeX/.bbl cleanup: each \\bibitem becomes its own
    paragraph and formatting commands are stripped, so the downstream splitter
    sees one clean reference per block."""
    s = _TEX_COMMENT_RE.sub("", s)

    m = re.search(r"\\begin\{thebibliography\}(?:\{[^}]*\})?(.*?)\\end\{thebibliography\}",
                  s, flags=re.S)
    if m:
        s = m.group(1)

    chunks = _BIBITEM_RE.split(s)
    if len(chunks) > 1:
        s = "\n\n".join(c.strip() for c in chunks[1:] if c.strip())

    s = s.replace("\\newblock", " ")
    s = _TEX_HREF_RE.sub(r"\2 (\1)", s)
    for _ in range(2):  # unwrap nested formatting like \emph{\textbf{...}}
        s = _TEX_WRAP_RE.sub(r"\1", s)
    s = _TEX_CMD_RE.sub(" ", s)
    s = s.replace("---", "—").replace("--", "–")
    s = s.replace("~", " ").replace("{", "").replace("}", "")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _extract_docx_sync(raw: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # References are frequently laid out in tables — plain paragraph
    # iteration silently drops them.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _extract_pdf_sync(raw: bytes) -> str:
    # Page cap bounds pathological PDFs; pdfminer is CPU-bound.
    return pdf_extract_text(io.BytesIO(raw), maxpages=MAX_PDF_PAGES) or ""


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

async def extract_files(files: List[UploadFile]) -> str:
    """Validate and extract text from uploads; returns concatenated text."""
    if not files:
        raise HTTPException(400, "No files uploaded.")
    if len(files) > MAX_FILES:
        raise HTTPException(400, f"Too many files (max {MAX_FILES}).")

    chunks: List[str] = []
    total_chars = 0

    for up in files:
        name = safe_name(up.filename)
        ext = _ext(name)

        if ext == ".doc":
            raise HTTPException(
                400, f"'{name}' is .doc (legacy). Please convert to .docx and re-upload.")
        if ext not in ALLOWED_EXTS:
            raise HTTPException(
                400, f"Unsupported file type for '{name}'. Allowed: {', '.join(ALLOWED_EXTS)}")

        raw = await read_limited(up, name)
        if not raw:
            raise HTTPException(400, f"'{name}' is empty.")
        validate_magic(ext, raw, name)

        t0 = time.monotonic()
        if ext in (".txt", ".tex", ".bbl"):
            text = _decode_text(raw)
            if ext in (".tex", ".bbl"):
                text = tex_to_text(text)
        elif ext == ".docx":
            try:
                text = await asyncio.to_thread(_extract_docx_sync, raw)
            except Exception:
                logger.exception("Failed to read DOCX: %s", name)
                raise HTTPException(400, f"Failed to read DOCX '{name}'. Is it a valid Word file?")
        else:  # .pdf
            if not PDF_AVAILABLE:
                raise HTTPException(
                    400, "PDF support is not available on the server. Install pdfminer.six.")
            try:
                text = await asyncio.to_thread(_extract_pdf_sync, raw)
            except Exception:
                logger.exception("Failed to read PDF: %s", name)
                raise HTTPException(400, f"Failed to read PDF '{name}'. Is it a valid PDF?")

        logger.info("Extracted '%s': %d bytes -> %d chars in %.2fs",
                    name, len(raw), len(text), time.monotonic() - t0)

        total_chars += len(text)
        if total_chars > MAX_TOTAL_TEXT_CHARS:
            raise HTTPException(
                413, "Uploads contain more text than can be processed in one request. "
                     "Split them into smaller batches.")
        chunks.append(text)

    combined = "\n\n".join(c.strip() for c in chunks if c and c.strip())
    if not combined:
        raise HTTPException(400, "No text could be extracted from the uploaded file(s).")
    return combined
