import os
from pathlib import Path

from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from refassist.graphs import run_one
from refassist.config import PipelineConfig
from docx import Document as DocxDocument
from typing import Optional, List, Tuple
import re
import asyncio
import io
import time
import zipfile
import logging

try:  # package-style run: `uvicorn api.app:app` from the repo root
    from api.uploads import extract_files, env_pos_int, MAX_FILES, MAX_UPLOAD_BYTES
    from api.jobs import STORE as JOB_STORE, Job
except ImportError:  # direct run from inside api/: `uvicorn app:app`
    from uploads import extract_files, env_pos_int, MAX_FILES, MAX_UPLOAD_BYTES
    from jobs import STORE as JOB_STORE, Job

# ---------- Setup ----------
BASE_DIR = Path(__file__).resolve().parent.parent  # repository root

from contextlib import asynccontextmanager


@asynccontextmanager
async def _lifespan(app):
    # Warm the Retraction Watch dataset in the background so the first
    # references already benefit from it (graceful no-op when offline).
    if os.getenv("REFASSIST_RW_PRELOAD", "1").lower() in ("1", "true", "yes"):
        from refassist.tools.retractionwatch import get_rw_db
        get_rw_db().start_background_load()
    yield


app = FastAPI(title="RefAssist API", version="0.8.0", lifespan=_lifespan)
app.mount("/new_ui", StaticFiles(directory=str(BASE_DIR / "web")), name="new_ui")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("refassist")

# One config for the process; env is read once at startup.
CFG = PipelineConfig()

# ---------- Request limits ----------
MAX_REFERENCES = env_pos_int("REFASSIST_MAX_REFERENCES", 200)

# Synchronous endpoints only process small batches inline; anything larger is
# auto-queued as a job (202) so no HTTP request can outlive a proxy timeout.
SYNC_MAX_REFS = env_pos_int("REFASSIST_SYNC_MAX_REFS", 10)

# Each reference runs a full multi-source + LLM pipeline; without a bound,
# one large upload would fan out into hundreds of concurrent pipelines.
_PIPELINE_SEM = asyncio.Semaphore(env_pos_int("REFASSIST_MAX_PARALLEL_REFS", 6))

# Wall-clock cap per reference: HTTP and per-source timeouts exist, but nothing
# else bounds a pathological correction loop from holding a semaphore slot.
REF_TIMEOUT_S = env_pos_int("REFASSIST_REF_TIMEOUT", 120, minimum=15)

# Cross-request result cache: users routinely fix one entry and re-check the
# whole list — every unchanged reference becomes a sub-millisecond cache hit
# instead of a full pipeline run (LLM calls included).
from cachetools import TTLCache
_RESULT_CACHE: TTLCache = TTLCache(
    maxsize=env_pos_int("REFASSIST_RESULT_CACHE_SIZE", 1000),
    ttl=env_pos_int("REFASSIST_RESULT_CACHE_TTL", 3600, minimum=0),
)


def _cache_key(ref: str) -> str:
    return re.sub(r"\s+", " ", ref).strip()

# Reject oversized requests from the Content-Length header, before any body
# parsing; the per-file chunked reads remain the hard backstop for clients
# that lie about or omit the header.
MAX_REQUEST_BYTES = MAX_FILES * MAX_UPLOAD_BYTES + 1024 * 1024


@app.middleware("http")
async def _reject_oversized_requests(request: Request, call_next):
    cl = request.headers.get("content-length")
    if cl and cl.isdigit() and int(cl) > MAX_REQUEST_BYTES:
        return JSONResponse(
            status_code=413,
            content={"detail": f"Request body exceeds {MAX_REQUEST_BYTES // (1024*1024)} MB."},
        )
    return await call_next(request)


# ---------- Per-IP rate limiting ----------
# Every processing request fans out into LLM calls billed to the operator, so
# expensive endpoints are rate limited per client IP (sliding window, in
# memory — per process, like the job store). Authentication for non-UI API
# exposure belongs in a reverse proxy in front of this app (see README).
from collections import deque

RATE_LIMIT_MAX = env_pos_int("REFASSIST_RATE_LIMIT_MAX", 30)      # requests…
RATE_LIMIT_WINDOW_S = env_pos_int("REFASSIST_RATE_LIMIT_WINDOW", 60)  # …per window
_RATE_BUCKETS: dict = {}
_RATE_LIMITED_PREFIXES = ("/api/jobs", "/api/process", "/api/download-report",
                          "/api/extract", "/v1/resolve", "/v1/upload")


def _rate_limited(ip: str) -> bool:
    now = time.time()
    bucket = _RATE_BUCKETS.get(ip)
    if bucket is None:
        bucket = _RATE_BUCKETS[ip] = deque()
    while bucket and now - bucket[0] > RATE_LIMIT_WINDOW_S:
        bucket.popleft()
    if len(bucket) >= RATE_LIMIT_MAX:
        return True
    bucket.append(now)
    if len(_RATE_BUCKETS) > 10_000:  # bound memory under IP churn
        _RATE_BUCKETS.clear()
    return False


@app.middleware("http")
async def _rate_limit(request: Request, call_next):
    if (request.method in ("POST", "DELETE")
            and request.url.path.startswith(_RATE_LIMITED_PREFIXES)):
        ip = request.client.host if request.client else "unknown"
        if _rate_limited(ip):
            return JSONResponse(
                status_code=429,
                content={"detail": "Too many requests — please slow down."},
                headers={"Retry-After": str(RATE_LIMIT_WINDOW_S)},
            )
    return await call_next(request)


async def _run_pipeline(ref: str) -> dict:
    key = _cache_key(ref)
    if (hit := _RESULT_CACHE.get(key)) is not None:
        logger.debug("Result cache hit for %r", key[:60])
        return hit
    async with _PIPELINE_SEM:
        # Double-check after waiting: an identical reference in the same batch
        # may have populated the cache while we held the queue.
        if (hit := _RESULT_CACHE.get(key)) is not None:
            return hit
        out = await asyncio.wait_for(run_one(ref, CFG), timeout=REF_TIMEOUT_S)
    _RESULT_CACHE[key] = out
    return out


# ---------- Smart reference splitting ----------
# Markers may be preceded by a stray quote (CSV-exported rows arrive as
# "[1] A. Author, ""Title,"" ..." — wrapped in quotes with internal quotes
# doubled).
_MARKER_PATTERNS = [
    re.compile(r"^\s*[\"']?\s*\[\d+\]"),  # [1]
    re.compile(r"^\s*[\"']?\s*\d+\."),    # 1.
    re.compile(r"^\s*[-•]"),              # - or •
]

# Signals that strongly suggest a boundary between references
_DOI_RE = re.compile(r"\b10\.\d{4,9}/\S+\b", re.I)
_URL_RE = re.compile(r"https?://\S+", re.I)
_YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
_AUTHOR_LIKE_RE = re.compile(r"^([A-Z][a-z]+|[A-Z]\.)[^\n]{0,60}\b([A-Z][a-z]+|[A-Z]\.)")  # crude author cue

def _has_any_marker(lines: List[str]) -> bool:
    for line in lines:
        if any(p.match(line) for p in _MARKER_PATTERNS):
            return True
    return False

def _strip_full_marker(line: str) -> str:
    return re.sub(r"^\s*[\"']?\s*(?:\[\d+\]|\d+\.|[-•])\s*", "", line)

def split_references(text: str) -> List[str]:
    """
    Smarter splitter:
      1) If markers like [1], '1.', '-' exist — use them.
      2) Else: split on blank lines (paragraphs).
      3) Else: heuristic segmentation using cues (title quotes/DOI/url/years/authors).
    """
    text = (text or "").strip()
    if not text:
        return []

    # Normalize
    text = re.sub(r"\r\n?", "\n", text)
    # CSV/Excel-escaped quotes: ""Title"" → "Title"
    text = re.sub(r'"{2,}', '"', text)
    # Condense multiple blank lines but remember boundaries
    text = re.sub(r"\n{3,}", "\n\n", text)

    lines = [ln.strip() for ln in text.split("\n")]

    # Path 1: marker-based
    if _has_any_marker(lines):
        refs: List[str] = []
        cur: List[str] = []
        for line in lines:
            if not line:
                continue
            if any(p.match(line) for p in _MARKER_PATTERNS):
                if cur:
                    refs.append(" ".join(cur).strip())
                cur = [_strip_full_marker(line)]
            else:
                cur.append(line)
        if cur:
            refs.append(" ".join(cur).strip())
        return [r for r in refs if r]

    # Path 2: paragraph-based (blank lines)
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(paragraphs) > 1:
        return paragraphs

    # Path 3: heuristic segmentation inside one big block
    blob = paragraphs[0] if paragraphs else text
    raw_lines = [ln.strip() for ln in blob.split("\n") if ln.strip()]

    # Build references by scanning for "strong start" lines
    refs: List[str] = []
    cur: List[str] = []

    def push_cur():
        if cur:
            refs.append(" ".join(cur).strip())

    def looks_like_new_ref(line: str) -> bool:
        # Starts like an author/title line, or contains strong identifiers (DOI/URL) in the previous chunk.
        if _AUTHOR_LIKE_RE.match(line):
            return True
        # Short lines with trailing period often begin titles in pasted refs
        if len(line) < 140 and line.endswith("."):
            return True
        # Container/venue cue
        if "Proc." in line or "Proceedings" in line or "IEEE" in line:
            return True
        return False

    for line in raw_lines:
        if not cur:
            cur.append(line)
            continue
        prev = " ".join(cur[-2:])  # last ~2 lines combined
        # boundary if current looks like a new ref AND previous had a DOI/URL/year
        prev_has_key = bool(_DOI_RE.search(prev) or _URL_RE.search(prev) or _YEAR_RE.search(prev))
        if looks_like_new_ref(line) and prev_has_key:
            push_cur()
            cur = [line]
        else:
            cur.append(line)

    push_cur()
    # Post-fix: if heuristics produced one giant ref, just return it
    return [r for r in refs if r] or [blob.strip()]


def _split_and_validate(refs_text: str) -> List[str]:
    refs = split_references(refs_text)
    if not refs:
        raise HTTPException(status_code=400, detail="No references detected in input")
    if len(refs) > MAX_REFERENCES:
        raise HTTPException(
            status_code=400,
            detail=f"Too many references ({len(refs)}). Maximum per request is {MAX_REFERENCES}."
        )
    return refs


# ---------- Formatting batch processing ----------
def _add_reference_section(doc: DocxDocument, entry: dict) -> None:
    """One professional section per reference in the .docx report."""
    doc.add_heading(f"Reference {entry['idx']}", level=2)
    data = entry.get("report_data") or {}

    if data:
        doc.add_paragraph(f"{data.get('status_label', '')} (confidence: {data.get('confidence', 'n/a')})")
        if data.get("retracted"):
            p = doc.add_paragraph()
            info = data.get("retraction_info") or {}
            note = ("RETRACTION NOTICE: this work has been retracted by its "
                    "publisher and should not be cited as valid research.")
            if info.get("reasons"):
                note += f" Reasons: {'; '.join(info['reasons'][:3])}."
            run = p.add_run(note)
            run.bold = True
        for a in data.get("action_items") or []:
            p = doc.add_paragraph(f"Action required: {a}", style="List Bullet")
            p.runs[0].bold = True
    doc.add_paragraph(f"Original: {entry['original']}")
    if entry.get("formatted"):
        doc.add_paragraph(f"Final (IEEE): {entry['formatted']}")

    corrections = (data.get("corrections") or []) if data else []
    if corrections:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Field", "Before", "After", "Source"
        for c in corrections:
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = (
                c["field"], c["old"], c["new"], c["source"])
    elif data:
        doc.add_paragraph("No corrections were needed.")

    warnings = (data.get("warnings") or []) if data else []
    for w in warnings:
        doc.add_paragraph(f"Warning: {w}", style="List Bullet")
    evidence = (data.get("evidence") or []) if data else []
    for e in evidence[:5]:
        doc.add_paragraph(f"{e['label']}: {e['url']}", style="List Bullet")
    if not data and entry.get("report"):
        doc.add_paragraph(entry["report"])


async def _process_one(idx: int, ref: str) -> dict:
    """Run the pipeline for one reference; never raises.

    Returns the full entry (including `report_data`, which the .docx builder
    needs); strip it with _public_entry() before returning JSON.
    """
    try:
        out = await _run_pipeline(ref.strip())
        formatted = out.get("formatted", ref)
        data = out.get("report_data") or {}
        return {
            "idx": idx, "original": ref, "formatted": formatted,
            "status": "success",
            "resolution": data.get("status"),
            "resolution_label": data.get("status_label"),
            "confidence": data.get("confidence"),
            "retracted": bool(data.get("retracted")),
            "retraction_info": data.get("retraction_info"),
            "author_mismatch": data.get("author_mismatch"),
            "fabrication": data.get("fabrication"),
            "corrections": data.get("corrections") or [],
            "warnings": data.get("warnings") or [],
            "bibtex": out.get("bibtex", "") or "",
            "doi": data.get("doi", "") or "",
            "report_data": data,
        }
    except Exception:
        logger.exception("Error processing reference %s", idx)
        return {
            "idx": idx, "original": ref, "formatted": ref + " [ERROR]",
            "status": "error", "resolution": "error",
            "resolution_label": "Error: reference could not be processed.",
            "confidence": "n/a", "retracted": False,
            "corrections": [], "warnings": [], "bibtex": "", "doi": "",
            "report_data": None,
        }


def _public_entry(entry: dict) -> dict:
    return {k: v for k, v in entry.items() if k != "report_data"}


def _annotate_duplicates(entries: List[dict]) -> None:
    """Two entries resolving to the same DOI is a classic bibliography error.
    Operates on copies (public entries), never on stored job results."""
    by_doi: dict = {}
    for e in entries:
        doi = (e.get("doi") or "").lower()
        if not doi:
            continue
        if doi in by_doi:
            first = by_doi[doi]
            e["duplicate_of"] = first["idx"]
            e["warnings"] = list(e["warnings"]) + [
                f"Duplicate: resolves to the same DOI as reference [{first['idx']}]."]
            first["warnings"] = list(first["warnings"]) + [
                f"Duplicate: reference [{e['idx']}] resolves to the same DOI."]
        else:
            by_doi[doi] = e


def _batch_response(detailed: List[dict]) -> dict:
    """JSON body shared by /api/process and completed jobs."""
    detailed = [_public_entry(e) for e in sorted(detailed, key=lambda e: e["idx"])]
    _annotate_duplicates(detailed)

    # Retracted works are marked in the output list itself — a user copying
    # the list must not be able to miss it.
    formatted_output = "\n".join(
        f"[{e['idx']}] {e['formatted']}" + ("  ⚠ [RETRACTED]" if e["retracted"] else "")
        for e in detailed
    )

    total_refs = len(detailed)
    success_count = sum(1 for r in detailed if r["status"] == "success")
    error_count = total_refs - success_count
    retracted_count = sum(1 for r in detailed if r["retracted"])

    preview_lines = [
        "Reference Processing Report", "=" * 30, "",
        f"Total references processed: {total_refs}",
        f"Successfully processed: {success_count}",
        f"Errors encountered: {error_count}",
    ]
    if retracted_count:
        preview_lines.append(f"RETRACTED works detected: {retracted_count}")
    preview_lines.append("")

    for entry in detailed[:10]:
        preview_lines.append(f"Reference {entry['idx']}: {entry['resolution_label'] or entry['status']}")
        if entry["retracted"]:
            preview_lines.append("  *** RETRACTED — do not cite without noting the retraction ***")
        preview_lines.append(f"  Original : {entry['original']}")
        if entry["status"] == "success":
            preview_lines.append(f"  Formatted: {entry['formatted']}")
        for c in entry["corrections"]:
            preview_lines.append(f"  Corrected {c['field']}: {c['old']} -> {c['new']}  [{c['source']}]")
        for w in entry["warnings"]:
            preview_lines.append(f"  Warning: {w}")
        preview_lines.append("")

    return {
        "success": True,
        "total_references": total_refs,
        "formatted_output": formatted_output,
        "preview": "\n".join(preview_lines),
        "results": detailed,
        "summary": {
            "total": total_refs, "success": success_count,
            "errors": error_count, "retracted": retracted_count,
            "duplicates": sum(1 for e in detailed if e.get("duplicate_of")),
        },
    }


def _build_report_doc(detailed: List[dict]) -> DocxDocument:
    detailed = sorted(detailed, key=lambda e: e["idx"])
    doc = DocxDocument()
    doc.add_heading("RefAssist Reference Processing Report", 0)
    ok = sum(1 for e in detailed if e.get("report_data"))
    doc.add_paragraph(f"{len(detailed)} reference(s) processed; {ok} completed successfully.")
    for entry in detailed:
        _add_reference_section(doc, entry)
    return doc


async def _process_batch(refs: List[str]) -> List[dict]:
    return list(await asyncio.gather(*(
        _process_one(i + 1, ref) for i, ref in enumerate(refs)
    )))


def _zip_results(formatted_refs: List[str], report_doc: DocxDocument,
                 detailed: Optional[List[dict]] = None) -> io.BytesIO:
    """Report bundle: formatted list, .docx report, .bib file, machine-readable JSON."""
    import json as _json
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        txt_content = "\n".join(f"[{i+1}] {ref}" for i, ref in enumerate(formatted_refs))
        zipf.writestr("formatted_references.txt", txt_content)

        docx_buffer = io.BytesIO()
        report_doc.save(docx_buffer)
        docx_buffer.seek(0)
        zipf.writestr("report.docx", docx_buffer.read())

        if detailed:
            entries = [_public_entry(e) for e in sorted(detailed, key=lambda e: e["idx"])]
            bib = "\n\n".join(e["bibtex"] for e in entries if e.get("bibtex"))
            if bib:
                zipf.writestr("references.bib", bib + "\n")
            zipf.writestr("results.json", _json.dumps(entries, indent=1, ensure_ascii=False))

    zip_buffer.seek(0)
    return zip_buffer


# ---------- Routes ----------
@app.get("/healthz")
async def healthz():
    return {"status": "ok"}


@app.get("/")
async def get_new_home(request: Request):
    return FileResponse(str(BASE_DIR / "web" / "index.html"))


# Single-reference resolver
class ResolveRequest(BaseModel):
    reference: str

@app.post("/v1/resolve")
async def resolve(req: ResolveRequest):
    if not req.reference.strip():
        raise HTTPException(status_code=400, detail="reference is required")
    try:
        out = await _run_pipeline(req.reference)
        return {
            "type": out.get("type"),
            "formatted": out.get("formatted"),
            "retracted": bool(out.get("retracted")),
            "fabrication": (out.get("report_data") or {}).get("fabrication"),
            "report": out.get("report"),
            "report_data": out.get("report_data"),
            "verification": out.get("verification"),
            "csl_json": out.get("csl_json"),
            "bibtex": out.get("bibtex"),
        }
    except Exception:
        logger.exception("resolve failed")
        raise HTTPException(status_code=500, detail="Internal error while resolving reference.")


# Server-side text extraction for uploaded files (multiple)
@app.post("/api/extract")
async def extract_files_endpoint(files: List[UploadFile] = File(...)):
    """
    Accepts multiple files and returns concatenated plain text suitable for splitting/processing.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")
    text = await extract_files(files)
    return {"text": text}



# Process pasted text OR (optionally) files: return JSON summary
@app.post("/api/process")
async def process_references_api(
    references: Optional[str] = Form(None),
    files: Optional[List[UploadFile]] = File(None)
):
    refs_text: str = (references or "").strip()

    # If files were sent directly to this endpoint, extract them here too
    if (not refs_text) and files:
        refs_text = (await extract_files(files)).strip()

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    refs = _split_and_validate(refs_text)
    if (queued := _queue_if_large(refs)) is not None:
        return queued
    detailed = await _process_batch(refs)
    return _batch_response(detailed)


def _start_job(refs: List[str]) -> Job:
    try:
        job = JOB_STORE.create(refs)
    except RuntimeError as e:
        raise HTTPException(status_code=429, detail=str(e))
    job.task = asyncio.create_task(_run_job(job))
    return job


def _queue_if_large(refs: List[str]) -> Optional[JSONResponse]:
    """Large batches must never run inside one HTTP request — auto-queue them."""
    if len(refs) <= SYNC_MAX_REFS:
        return None
    job = _start_job(refs)
    return JSONResponse(status_code=202, content={
        "success": False,
        "queued": True,
        "job_id": job.id,
        "total": job.total,
        "status_url": f"/api/jobs/{job.id}",
        "report_url": f"/api/jobs/{job.id}/report",
        "detail": (f"Batch of {len(refs)} references exceeds the synchronous "
                   f"limit ({SYNC_MAX_REFS}); poll status_url for results."),
    })


# ---------- Bulk processing jobs ----------
# Large batches run for many minutes; jobs decouple submission from retrieval
# so no HTTP request ever has to outlive a proxy timeout.

async def _run_job(job: Job) -> None:
    job.status = "running"
    try:
        async def worker(idx: int, ref: str) -> None:
            entry = await _process_one(idx, ref)
            job.results.append(entry)
            job.done += 1

        await asyncio.gather(*(worker(i + 1, r) for i, r in enumerate(job.refs)))
        job.status = "completed"
    except asyncio.CancelledError:
        job.status = "cancelled"
    except Exception:
        logger.exception("Job %s failed", job.id)
        job.status = "failed"
        job.error = "Internal error while processing the batch."
    finally:
        job.finished_at = time.time()
        logger.info("Job %s %s (%d/%d references)", job.id, job.status, job.done, job.total)


def _job_payload(job: Job) -> dict:
    payload = {
        "job_id": job.id,
        "status": job.status,
        "total": job.total,
        "done": job.done,
    }
    if job.error:
        payload["error"] = job.error
    if job.status == "completed":
        payload.update(_batch_response(job.results))
    elif job.results:
        # Progressive results so the UI can render entries as they finish.
        payload["results"] = [
            _public_entry(e) for e in sorted(job.results, key=lambda e: e["idx"])
        ]
    return payload


@app.post("/api/jobs")
async def create_job(
    references: Optional[str] = Form(None),
    files: Optional[List[UploadFile]] = File(None)
):
    refs_text: str = (references or "").strip()
    if (not refs_text) and files:
        refs_text = (await extract_files(files)).strip()
    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    refs = _split_and_validate(refs_text)
    job = _start_job(refs)
    return {
        "job_id": job.id,
        "status": job.status,
        "total": job.total,
        "status_url": f"/api/jobs/{job.id}",
        "report_url": f"/api/jobs/{job.id}/report",
    }


@app.get("/api/jobs/{job_id}")
async def job_status(job_id: str):
    job = JOB_STORE.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Unknown or expired job.")
    return _job_payload(job)


@app.delete("/api/jobs/{job_id}")
async def cancel_job(job_id: str):
    job = JOB_STORE.cancel(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Unknown or expired job.")
    return {"job_id": job.id, "status": job.status}


@app.get("/api/jobs/{job_id}/report")
async def job_report(job_id: str):
    """ZIP report built from STORED results — the pipeline is never re-run."""
    job = JOB_STORE.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Unknown or expired job.")
    if job.status != "completed":
        raise HTTPException(status_code=409, detail=f"Job is {job.status}, not completed.")

    detailed = sorted(job.results, key=lambda e: e["idx"])
    formatted_refs = [
        e["formatted"] + ("  [RETRACTED]" if e["retracted"] else "") for e in detailed
    ]
    return StreamingResponse(
        _zip_results(formatted_refs, _build_report_doc(detailed), detailed),
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="refassist_report.zip"'}
    )


# Synchronous report ZIP (small batches; re-processes the input)
@app.post("/api/download-report")
async def download_full_report(
    references: Optional[str] = Form(None),
    files: Optional[List[UploadFile]] = File(None)
):
    refs_text: str = (references or "").strip()
    if (not refs_text) and files:
        refs_text = (await extract_files(files)).strip()

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    refs = _split_and_validate(refs_text)
    if (queued := _queue_if_large(refs)) is not None:
        return queued
    detailed = await _process_batch(refs)
    formatted_refs = [
        e["formatted"] + ("  [RETRACTED]" if e["retracted"] else "")
        for e in sorted(detailed, key=lambda e: e["idx"])
    ]
    return StreamingResponse(
        _zip_results(formatted_refs, _build_report_doc(detailed), detailed),
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="refassist_report.zip"'}
    )


# Legacy batch (kept)
@app.post("/v1/upload")
async def upload_references_legacy(
    references: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    refs_text: Optional[str] = (references or "").strip()
    if (not refs_text) and file:
        refs_text = (await extract_files([file])).strip()

    if not refs_text:
        raise HTTPException(status_code=400, detail="No references provided")

    refs = _split_and_validate(refs_text)
    if (queued := _queue_if_large(refs)) is not None:
        return queued
    detailed = await _process_batch(refs)
    formatted_refs = [e["formatted"] for e in sorted(detailed, key=lambda e: e["idx"])]
    return StreamingResponse(
        _zip_results(formatted_refs, _build_report_doc(detailed), detailed),
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="results.zip"'}
    )
